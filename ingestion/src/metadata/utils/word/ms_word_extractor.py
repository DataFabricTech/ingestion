# Copyright 2024 Mobigen
# Licensed under the Apache License, Version 2.0 (the "License")
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
#
# Notice!
# This software is based on https://open-metadata.org and has been modified accordingly.

import os

from docx import Document
from tika import parser

from metadata.ml.summarization import Summarization


class MsWordMetadataExtractor:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.summarizer = Summarization()

    def extract_metadata(self) -> dict:

        # Tika로 문서 파싱
        parsed = parser.from_file(self.file_path)
        # 메타데이터 추출
        metadata = parsed['metadata']
        # 파일 확장자가 docx가 아니면 tika 결과만을 반환
        file_extension = os.path.splitext(self.file_path)[1]
        if file_extension != '.docx':
            return metadata

        # 문서 열기
        doc = Document(self.file_path)

        # 메타데이터 추출
        core_properties = doc.core_properties
        if core_properties.author is not None:
            metadata['Author'] = core_properties.author
        if core_properties.category is not None:
            metadata['Category'] = core_properties.category
        if core_properties.comments is not None:
            metadata['Comments'] = core_properties.comments
        if core_properties.content_status is not None:
            metadata['Content Status'] = core_properties.content_status
        if core_properties.created is not None:
            metadata['Created'] = core_properties.created
        if core_properties.identifier is not None:
            metadata['Identifier'] = core_properties.identifier
        if core_properties.keywords is not None:
            metadata['Keywords'] = core_properties.keywords
        if core_properties.language is not None:
            metadata['Language'] = core_properties.language
        if core_properties.last_modified_by is not None:
            metadata['Last Modified By'] = core_properties.last_modified_by
        if core_properties.modified is not None:
            metadata['Modified'] = core_properties.modified
        if core_properties.revision is not None:
            metadata['Revision'] = core_properties.revision
        if core_properties.subject is not None:
            metadata['Subject'] = core_properties.subject
        if core_properties.title is not None:
            metadata['Title'] = core_properties.title
        if core_properties.version is not None:
            metadata['Version'] = core_properties.version

        sample_data = self.get_sample_data(-1)
        if sample_data is not None:
            str_summary = self.summarizer.summarize(sample_data)
            metadata['Summary'] = str_summary
        return metadata

    def get_sample_data(self, chunk_size: int = 1000) -> str:
        sample_text = ""
        # Word 문서를 불러 옵니다.
        doc = Document(self.file_path)
        # 문서의 모든 단락을 순회
        for para in doc.paragraphs:
            # 단락이 비어 있지 않으면 추가
            if para.text.strip():
                sample_text += para.text
                sample_text += "\n"
                if 0 < chunk_size < len(sample_text):
                    break
        return sample_text


# # 3. Word 문서로부터 텍스트를 추출하고 주제를 예측하는 함수
# def predict_topic(docx_file, model):
#     # 문서로부터 텍스트 추출
#     document_text = extract_text_from_docx(docx_file)
#
#     # 예측
#     predicted_topic = model.predict([document_text])
#
#     return predicted_topic[0]
#
# # 4. 사용 예시
# if __name__ == "__main__":
#     # Word 파일 경로
#     docx_file_path = 'example.docx'
#
#     # 주제 분류기 생성 및 학습
#     topic_classifier = create_topic_classifier()
#
#     # 문서의 주제 예측
#     topic = predict_topic(docx_file_path, topic_classifier)
#     print(f"The predicted topic of the document is: {topic}")
#
# from docx import Document
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.naive_bayes import MultinomialNB
# from sklearn.pipeline import make_pipeline
# from konlpy.tag import Okt
# import nltk
#
# # NLTK의 punkt 패키지를 다운로드 (첫 실행 시 필요)
# nltk.download('punkt')
#
# # 1. MS Word 파일에서 텍스트 추출
# def extract_text_from_docx(docx_file):
#     doc = Document(docx_file)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     return '\n'.join(full_text)
#
# # 2. 한국어 형태소 분석 및 토큰화
# def tokenize_korean(text):
#     okt = Okt()
#     tokens = okt.morphs(text)  # 형태소 분석 및 토큰화
#     return ' '.join(tokens)
#
#
# # 3. 한국어 문서 주제 분류 모델 생성
# def create_topic_classifier():
#     # 예시 문서 데이터 (한국어)
#     documents = [
#         "이 문서는 데이터 과학에 대해 설명하고 있습니다.",
#         "3분기 재무 보고서가 제출되었습니다.",
#         "이 문서는 두 당사자 간의 법적 계약입니다.",
#         "이 연구 논문은 인공지능의 최신 발전을 다룹니다.",
#         "회사는 연간 예산을 준비하고 있습니다.",
#         "최근 인공지능과 머신러닝의 발전은 매우 흥미롭습니다.",
#         "이 계약서는 매수인과 매도인 간의 합의입니다.",
#         "최신 분기 실적 보고서가 발행되었습니다."
#     ]
#
#     # 각 문서에 해당하는 주제 레이블 (한국어)
#     labels = [
#         '데이터 과학',
#         '재무',
#         '법률',
#         '인공지능',
#         '재무',
#         '인공지능',
#         '법률',
#         '재무'
#     ]
#
#     # 한국어 문서를 형태소 분석 및 토큰화
#     tokenized_documents = [tokenize_korean(doc) for doc in documents]
#
#     # TF-IDF와 나이브 베이즈 분류기를 사용하는 파이프라인 생성
#     model = make_pipeline(TfidfVectorizer(), MultinomialNB())
#
#     # 모델 학습
#     model.fit(tokenized_documents, labels)
#
#     return model
#
#
# # 4. Word 문서로부터 텍스트를 추출하고 주제를 예측하는 함수
# def predict_topic(docx_file, model):
#     # 문서로부터 텍스트 추출
#     document_text = extract_text_from_docx(docx_file)
#
#     # 문서 텍스트를 한국어 형태소 분석 및 토큰화
#     tokenized_text = tokenize_korean(document_text)
#
#     # 예측
#     predicted_topic = model.predict([tokenized_text])
#
#     return predicted_topic[0]
#
#
# # 5. 사용 예시
# if __name__ == "__main__":
#     # Word 파일 경로
#     docx_file_path = 'example.docx'
#
#     # 주제 분류기 생성 및 학습
#     topic_classifier = create_topic_classifier()
#
#     # 문서의 주제 예측
#     topic = predict_topic(docx_file_path, topic_classifier)
#     print(f"The predicted topic of the document is: {topic}")
